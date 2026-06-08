#!/usr/bin/env python3
"""Generate CCL Workflow Logic Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# --- Styles helpers ---
def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Calibri'
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '6')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get(edge, 'BFBFBF'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CCL Questionnaire – Workflow Logic')
set_font(run, size=20, bold=True, color=(31, 73, 125))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('NBG RDARR Compliance Checklist Application')
set_font(run2, size=12, italic=True, color=(89, 89, 89))

doc.add_paragraph()

# ============================================================
# 1. OVERVIEW
# ============================================================
add_heading(doc, '1. Overview', level=1)

overview_text = (
    'The CCL Questionnaire is a multi-stage compliance workflow used to collect, review, '
    'and validate Business Unit (BU) responses to BCBS 239 / ECB RDARR requirements across '
    'annual questionnaire cycles. The application enforces strict role-based access control '
    'and a sequential state machine for both Cycles and individual Validations. '
    'Every significant action is recorded in a persistent audit log.'
)
p = doc.add_paragraph(overview_text)
p.runs[0].font.name = 'Calibri'
p.runs[0].font.size = Pt(11)

doc.add_paragraph()

# ============================================================
# 2. ROLES
# ============================================================
add_heading(doc, '2. User Roles', level=1)

roles = [
    ('Admin', 'Creates cycles, closes cycles, manages users and applicability. Supersedes Validator permissions.'),
    ('Validator', 'Configures question applicability, submits cycles for approval, distributes cycles, updates and submits individual validations for senior review, and may return BU responses.'),
    ('Senior Validator', 'Approves or rejects cycles (draft ↔ published) and approves or rejects individual validation assessments (pending_approval → closed / in_review).'),
    ('Responder', 'Belongs to one or more Business Units. Saves and submits responses to assigned questions.'),
    ('Viewer', 'Read-only access to cycles, responses, and validation data.'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.LEFT

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Role'
hdr_cells[1].text = 'Responsibilities'
for cell in hdr_cells:
    shade_cell(cell, '1F497D')
    for para in cell.paragraphs:
        for run in para.runs:
            set_font(run, bold=True, color=(255, 255, 255))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for role, desc in roles:
    row = table.add_row()
    row.cells[0].text = role
    row.cells[1].text = desc
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(4.5)
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)

doc.add_paragraph()

# ============================================================
# 3. CYCLE WORKFLOW
# ============================================================
add_heading(doc, '3. Questionnaire Cycle Workflow', level=1)

intro = doc.add_paragraph(
    'A Questionnaire Cycle represents one annual assessment period. It moves through five statuses '
    'in a fixed sequence. Only the transitions listed below are permitted by the system.'
)
intro.runs[0].font.name = 'Calibri'
intro.runs[0].font.size = Pt(11)

doc.add_paragraph()

# Cycle state table
cycle_states = [
    ('draft', 'Cycle created; applicability and questions configured. Validator may edit.', 'Admin (create)\nValidator (edit applicability)'),
    ('pending_approval', 'Submitted to Senior Validator for approval. Read-only to all others.', 'Validator (submit)'),
    ('published', 'Approved by Senior Validator; ready to be distributed to BUs.', 'Senior Validator (approve)'),
    ('distributed', 'Active — BU responses are being collected. Responses visible and editable by Responders.', 'Validator (distribute)'),
    ('closed', 'Collection complete; all validation work is finalised.', 'Admin (close)'),
]

add_heading(doc, '3.1 Cycle Statuses', level=2)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'

for i, hdr in enumerate(['Status', 'Meaning', 'Who enters this status']):
    cell = tbl.rows[0].cells[i]
    cell.text = hdr
    shade_cell(cell, '2E75B6')
    for para in cell.paragraphs:
        for run in para.runs:
            set_font(run, bold=True, color=(255, 255, 255))

for status, meaning, who in cycle_states:
    row = tbl.add_row()
    row.cells[0].text = status
    row.cells[1].text = meaning
    row.cells[2].text = who
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)

doc.add_paragraph()

add_heading(doc, '3.2 Cycle State Transitions', level=2)

cycle_transitions = [
    ('draft → pending_approval', 'Validator', 'PUT /api/cycles/:id/submit', 'Cycle must be in draft. Clears any prior rejection comment.'),
    ('pending_approval → published', 'Senior Validator', 'PUT /api/cycles/:id/approve', 'Cycle must be pending_approval. Sets published_at timestamp.'),
    ('pending_approval → draft', 'Senior Validator', 'PUT /api/cycles/:id/reject', 'A rejection comment is mandatory. Sets rejection_comment on the cycle for Validator to read.'),
    ('published → distributed', 'Validator', 'PUT /api/cycles/:id/distribute', 'Cycle must be published. Triggers BU response collection. Sets distributed_at.'),
    ('distributed → closed', 'Admin', 'PUT /api/cycles/:id/close', 'Cycle must be distributed. Sets closed_at.'),
]

tbl2 = doc.add_table(rows=1, cols=4)
tbl2.style = 'Table Grid'

for i, hdr in enumerate(['Transition', 'Actor', 'API Endpoint', 'Guard / Notes']):
    cell = tbl2.rows[0].cells[i]
    cell.text = hdr
    shade_cell(cell, '2E75B6')
    for para in cell.paragraphs:
        for run in para.runs:
            set_font(run, bold=True, color=(255, 255, 255))

for transition, actor, endpoint, notes in cycle_transitions:
    row = tbl2.add_row()
    row.cells[0].text = transition
    row.cells[1].text = actor
    row.cells[2].text = endpoint
    row.cells[3].text = notes
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

doc.add_paragraph()

# ============================================================
# 4. RESPONSE WORKFLOW
# ============================================================
add_heading(doc, '4. BU Response Workflow', level=1)

resp_intro = doc.add_paragraph(
    'When a cycle reaches distributed status, one response record is created per (Business Unit, Question) '
    'combination as defined in the question_applicability table. Responders — assigned to one or more '
    'BU codes — fill in their compliance score (1–4) and comments, then submit.'
)
resp_intro.runs[0].font.name = 'Calibri'
resp_intro.runs[0].font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, '4.1 Response Statuses', level=2)

resp_states = [
    ('draft', 'Initial state. Response exists but has not been touched by a Responder.'),
    ('in_progress', 'Responder has saved work at least once but has not yet submitted.'),
    ('submitted', 'Responder has finalised and submitted the response. Locked for Responder edits.'),
]

tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = 'Table Grid'

for i, hdr in enumerate(['Status', 'Meaning']):
    cell = tbl3.rows[0].cells[i]
    cell.text = hdr
    shade_cell(cell, '2E75B6')
    for para in cell.paragraphs:
        for run in para.runs:
            set_font(run, bold=True, color=(255, 255, 255))

for status, meaning in resp_states:
    row = tbl3.add_row()
    row.cells[0].text = status
    row.cells[1].text = meaning
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)

doc.add_paragraph()
add_heading(doc, '4.2 Response Transitions', level=2)

resp_transitions = [
    ('draft → in_progress', 'Responder', 'PUT /api/cycles/:cycleId/responses/:id', 'Triggered automatically on first save (score or comment).'),
    ('in_progress → submitted', 'Responder', 'PUT /api/cycles/:cycleId/responses/:id/submit', 'Sets submitted_at. Triggers validation upsert (see §5).'),
    ('submitted → in_progress', 'Validator / Admin', 'PUT /api/cycles/:cycleId/responses/:id/return', 'Returns the response for rework. A return_comment may be provided. Also resets the parent validation to pending (see §5.3).'),
]

tbl4 = doc.add_table(rows=1, cols=4)
tbl4.style = 'Table Grid'

for i, hdr in enumerate(['Transition', 'Actor', 'API Endpoint', 'Notes']):
    cell = tbl4.rows[0].cells[i]
    cell.text = hdr
    shade_cell(cell, '2E75B6')
    for para in cell.paragraphs:
        for run in para.runs:
            set_font(run, bold=True, color=(255, 255, 255))

for transition, actor, endpoint, notes in resp_transitions:
    row = tbl4.add_row()
    row.cells[0].text = transition
    row.cells[1].text = actor
    row.cells[2].text = endpoint
    row.cells[3].text = notes
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

doc.add_paragraph()

# Trigger note
trigger = doc.add_paragraph()
run = trigger.add_run('Automatic trigger: ')
set_font(run, bold=True)
run2 = trigger.add_run(
    'When the last pending BU response for a question is submitted (i.e. ALL response '
    'records for that cycle+question are in submitted status), the system automatically '
    'upserts a validation record with status in_review, making it available for Validator review.'
)
set_font(run2)

doc.add_paragraph()

# ============================================================
# 5. VALIDATION WORKFLOW
# ============================================================
add_heading(doc, '5. Validation Workflow', level=1)

val_intro = doc.add_paragraph(
    'One Validation record exists per (Cycle, Question). It aggregates all BU responses for that question '
    'and captures the Validator\'s consolidated assessment. Validations go through a two-tier approval: '
    'first the Validator, then the Senior Validator.'
)
val_intro.runs[0].font.name = 'Calibri'
val_intro.runs[0].font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, '5.1 Validation Statuses', level=2)

val_states = [
    ('pending', 'Not all BU responses have been submitted yet. Validation cannot be acted upon.'),
    ('in_review', 'All BU responses are submitted. Validator can review responses, set a validation_score, write justification and additional_controls notes.'),
    ('pending_approval', 'Validator has submitted the assessment for Senior Validator sign-off.'),
    ('closed', 'Senior Validator has approved. Validation is finalised and locked.'),
]

tbl5 = doc.add_table(rows=1, cols=2)
tbl5.style = 'Table Grid'

for i, hdr in enumerate(['Status', 'Meaning']):
    cell = tbl5.rows[0].cells[i]
    cell.text = hdr
    shade_cell(cell, '2E75B6')
    for para in cell.paragraphs:
        for run in para.runs:
            set_font(run, bold=True, color=(255, 255, 255))

for status, meaning in val_states:
    row = tbl5.add_row()
    row.cells[0].text = status
    row.cells[1].text = meaning
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)

doc.add_paragraph()
add_heading(doc, '5.2 Validation Transitions', level=2)

val_transitions = [
    ('pending → in_review', 'System (automatic)', '—', 'Triggered when all BU responses for the question are submitted.'),
    ('in_review → pending_approval', 'Validator / Admin', 'PUT /api/cycles/:cycleId/validations/:id/close', 'Validator submits consolidated assessment for Senior Validator approval. Records submission in workflow_history.'),
    ('pending_approval → closed', 'Senior Validator / Admin', 'PUT /api/cycles/:cycleId/validations/:id/approve', 'Final approval. Sets senior_validated_by and senior_validated_at. Records approval in workflow_history.'),
    ('pending_approval → in_review', 'Senior Validator / Admin', 'PUT /api/cycles/:cycleId/validations/:id/reject', 'Sends back to Validator for revision. Sets senior_rejection_comment. Records rejection in workflow_history.'),
]

tbl6 = doc.add_table(rows=1, cols=4)
tbl6.style = 'Table Grid'

for i, hdr in enumerate(['Transition', 'Actor', 'API Endpoint', 'Notes']):
    cell = tbl6.rows[0].cells[i]
    cell.text = hdr
    shade_cell(cell, '2E75B6')
    for para in cell.paragraphs:
        for run in para.runs:
            set_font(run, bold=True, color=(255, 255, 255))

for transition, actor, endpoint, notes in val_transitions:
    row = tbl6.add_row()
    row.cells[0].text = transition
    row.cells[1].text = actor
    row.cells[2].text = endpoint
    row.cells[3].text = notes
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

doc.add_paragraph()

add_heading(doc, '5.3 Impact of Returning a Response on Validations', level=2)
ret_note = doc.add_paragraph(
    'When a Validator returns a submitted response to in_progress, the system also resets the parent '
    'validation record from in_review → pending (if it is currently in_review). This ensures the '
    'validation queue correctly reflects that not all BU responses are finalised for that question.'
)
ret_note.runs[0].font.name = 'Calibri'
ret_note.runs[0].font.size = Pt(11)

doc.add_paragraph()

# ============================================================
# 6. APPLICABILITY
# ============================================================
add_heading(doc, '6. Question Applicability', level=1)

app_text = doc.add_paragraph(
    'Before a cycle is distributed, Admins and Validators configure which Business Units must answer '
    'which questions via the question_applicability table. Each record links a (cycle_id, question_id, bu_code) '
    'triple. Applicability can only be modified while the cycle is in draft or published status; '
    'it is frozen once the cycle is distributed or closed.'
)
app_text.runs[0].font.name = 'Calibri'
app_text.runs[0].font.size = Pt(11)

doc.add_paragraph()

# ============================================================
# 7. END-TO-END FLOW
# ============================================================
add_heading(doc, '7. End-to-End Workflow Summary', level=1)

e2e = doc.add_paragraph()

steps = [
    ('1', 'Admin creates a new Cycle', 'Cycle status: draft'),
    ('2', 'Validator / Admin assigns BUs to questions', 'question_applicability records created'),
    ('3', 'Validator submits cycle for approval', 'Cycle: draft → pending_approval'),
    ('4a', 'Senior Validator approves cycle', 'Cycle: pending_approval → published'),
    ('4b', 'Senior Validator rejects cycle', 'Cycle: pending_approval → draft (with comment)'),
    ('5', 'Validator distributes the cycle', 'Cycle: published → distributed; response records initialised'),
    ('6', 'Responders save and submit BU responses', 'Response: draft → in_progress → submitted'),
    ('7', 'System detects all BU responses submitted for a question', 'Validation: pending → in_review (automatic)'),
    ('8', 'Validator reviews responses, sets score and justification', 'Validation stays in_review while being edited'),
    ('9', 'Validator submits validation for approval', 'Validation: in_review → pending_approval'),
    ('10a', 'Senior Validator approves validation', 'Validation: pending_approval → closed'),
    ('10b', 'Senior Validator rejects validation', 'Validation: pending_approval → in_review (with rejection_comment)'),
    ('11', 'Steps 8–10 repeat until all validations are closed', '—'),
    ('12', 'Admin closes the cycle', 'Cycle: distributed → closed'),
]

tbl7 = doc.add_table(rows=1, cols=3)
tbl7.style = 'Table Grid'

for i, hdr in enumerate(['Step', 'Action', 'State Change']):
    cell = tbl7.rows[0].cells[i]
    cell.text = hdr
    shade_cell(cell, '1F497D')
    for para in cell.paragraphs:
        for run in para.runs:
            set_font(run, bold=True, color=(255, 255, 255))

for step, action, state in steps:
    row = tbl7.add_row()
    row.cells[0].text = step
    row.cells[1].text = action
    row.cells[2].text = state
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(3.5)
    row.cells[2].width = Inches(2.5)
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)

doc.add_paragraph()

# ============================================================
# 8. AUDIT LOG
# ============================================================
add_heading(doc, '8. Audit Log', level=1)

audit_text = doc.add_paragraph(
    'Every workflow transition is recorded in the audit_log table and (for validations) appended to '
    'the validation\'s workflow_history JSONB array. Each entry captures: entity type, entity ID, '
    'action name, actor ID, actor display name, actor role, old value, new value, and timestamp. '
    'The following actions are audited:'
)
audit_text.runs[0].font.name = 'Calibri'
audit_text.runs[0].font.size = Pt(11)

audit_actions = [
    'cycle_created', 'cycle_submitted_for_approval', 'cycle_approved', 'cycle_rejected',
    'cycle_distributed', 'cycle_closed',
    'response_saved', 'response_submitted', 'response_returned',
    'validation_updated', 'validation_submitted_for_approval', 'validation_approved', 'validation_rejected',
    'applicability_assigned', 'applicability_removed',
]

for action in audit_actions:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(action)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

doc.add_paragraph()

# ============================================================
# 9. ATTACHMENTS
# ============================================================
add_heading(doc, '9. File Attachments', level=1)

att_text = doc.add_paragraph(
    'Both responses and validations support file attachments (response_attachments and '
    'validation_attachments tables). Files are stored on disk in the uploads/ directory and '
    'referenced by file_name and file_path. Attachments are deleted when the parent record is deleted.'
)
att_text.runs[0].font.name = 'Calibri'
att_text.runs[0].font.size = Pt(11)

doc.add_paragraph()

# ============================================================
# Footer note
# ============================================================
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('Generated automatically from the CCLQuestionnaire codebase · June 2026')
set_font(run, size=9, italic=True, color=(128, 128, 128))

# Save
output_path = '/home/esakelar/CCLQuestionnaire/CCL_Workflow_Logic.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
